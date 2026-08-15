#!/usr/bin/env python3
"""Splits the Paramatma-sandarbha source docx into separate physical IAST
docx files, before any transliteration -- see generate_paramatma_
sandarbha.py's module docstring for the full methodology (same as
Bhagavat-sandarbha's pipeline).

Outputs (still IAST, ready for convert_no_hyphen_join.py):
  paramatma_sandarbha_anucchedas.docx           -- mula content, per anuccheda
  paramatma_sandarbha_sarva_samvadini.docx      -- her voice, per anuccheda
  paramatma_sandarbha_sarva_samvadini_anuvyakhya.docx
      -- her further excerpt after mula's own closing colophon (see
         SARVA_SAMVADINI_ANUVYAKHYA_* in generate_paramatma_sandarbha.py)

Unlike Bhagavat-sandarbha, this docx has no separate "sāra-niṣkarṣa"
essence-summary and no "Appendix" section.
"""
import sys

from docx import Document

sys.path.insert(0, ".")
from docx_walk import iter_all_paragraphs
from paragraph_rewrite import paragraph_full_text
from generate_paramatma_sandarbha import (
    DOCX_PATH,
    SARVA_SAMVADINI_ANUVYAKHYA_START, SARVA_SAMVADINI_ANUVYAKHYA_END,
    _SARVA_PREFIX_RE,
    find_chunks,
)

ANUCCHEDAS_OUT = "paramatma_sandarbha_anucchedas.docx"
SARVA_SAMVADINI_OUT = "paramatma_sandarbha_sarva_samvadini.docx"
SARVA_SAMVADINI_ANUVYAKHYA_OUT = "paramatma_sandarbha_sarva_samvadini_anuvyakhya.docx"


def copy_paragraph(src_paragraph, dest_doc, override_text=None):
    """Copies a paragraph run-by-run, preserving text and color. If
    override_text is given (the "sarva-saṁvādinī:" label -- possibly with
    an inline "[N]"/"[N-M]" self-reference in front of it, see
    _SARVA_PREFIX_RE -- already stripped off), the paragraph is written as
    a single run with that text instead.
    """
    dest_p = dest_doc.add_paragraph()
    if override_text is not None:
        dest_p.add_run(override_text)
        return dest_p
    for run in src_paragraph.runs:
        new_run = dest_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.color is not None and run.font.color.type is not None:
            try:
                new_run.font.color.rgb = run.font.color.rgb
            except Exception:
                pass
    return dest_p


def extract_flat_range(orig_paras, orig_para_objs, start, end, out_path):
    """Copies a flat idx range (no "[N]" markers, no mula/sarva split) to
    its own docx -- used for the standalone trailing section (her further
    excerpt) that sits outside any anuccheda's own range."""
    doc = Document()
    n = 0
    for idx in range(start, end):
        t = orig_paras[idx].strip()
        if not t or "o)0(o" in t:
            continue
        copy_paragraph(orig_para_objs[idx], doc, None)
        n += 1
    doc.save(out_path)
    return n


def main():
    src_doc = Document(DOCX_PATH)
    orig_para_objs = list(iter_all_paragraphs(src_doc))
    orig_paras = [paragraph_full_text(p).strip() for p in orig_para_objs]

    chunks, marker_indices = find_chunks(orig_paras)

    anucchedas_doc = Document()
    sarva_doc = Document()

    n_anuccheda_paras = 0
    n_sarva_paras = 0
    n_anuccheda_with_sarva = 0

    for label, start, end in chunks:
        mula_items = []   # (idx, override_text_or_None)
        sarva_items = []
        seen_label = False
        for idx in range(start, end):
            if idx in marker_indices:
                continue  # decimal sub-part markers, no content of their own
            t = orig_paras[idx].strip()
            if not t or "o)0(o" in t:
                continue
            m = _SARVA_PREFIX_RE.match(t)
            if m:
                seen_label = True
            override = t[m.end():] if m else None
            (sarva_items if seen_label else mula_items).append((idx, override))

        if label is not None:
            anucchedas_doc.add_paragraph().add_run(f"[{label}]")
            n_anuccheda_paras += 1
        for idx, override in mula_items:
            copy_paragraph(orig_para_objs[idx], anucchedas_doc, override)
            n_anuccheda_paras += 1

        if sarva_items:
            if label is not None:
                sarva_doc.add_paragraph().add_run(f"[{label}]")
                n_sarva_paras += 1
            n_anuccheda_with_sarva += 1
            for idx, override in sarva_items:
                copy_paragraph(orig_para_objs[idx], sarva_doc, override)
                n_sarva_paras += 1

    anucchedas_doc.save(ANUCCHEDAS_OUT)
    sarva_doc.save(SARVA_SAMVADINI_OUT)
    print(f"{ANUCCHEDAS_OUT}: {n_anuccheda_paras} paragraphs")
    print(f"{SARVA_SAMVADINI_OUT}: {n_sarva_paras} paragraphs "
          f"({n_anuccheda_with_sarva} chunks with sarva-samvadini content)")

    n_anuvyakhya_paras = extract_flat_range(
        orig_paras, orig_para_objs, SARVA_SAMVADINI_ANUVYAKHYA_START,
        SARVA_SAMVADINI_ANUVYAKHYA_END, SARVA_SAMVADINI_ANUVYAKHYA_OUT
    )
    print(f"{SARVA_SAMVADINI_ANUVYAKHYA_OUT}: {n_anuvyakhya_paras} paragraphs")


if __name__ == "__main__":
    main()
