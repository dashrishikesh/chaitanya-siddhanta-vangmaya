"""Helpers to reach every paragraph in a .docx file that python-docx's
high-level API doesn't expose by default: table cells (nested), headers/
footers, and footnotes/endnotes (python-docx has no footnote API at all
as of 1.2.0 -- we reach into the underlying OOXML part ourselves).
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

FOOTNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
ENDNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"


def iter_table_paragraphs(table):
    """Yield every Paragraph in a table, including nested tables inside cells."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_body_paragraphs(document):
    """Body paragraphs + all table paragraphs, in document order-ish
    (python-docx doesn't give true interleaved order across tables and
    paragraphs trivially, so this yields all top-level paragraphs first,
    then all table paragraphs -- fine for a global convert pass since we
    process every paragraph regardless of order).
    """
    yield from document.paragraphs
    for table in document.tables:
        yield from iter_table_paragraphs(table)


def iter_header_footer_paragraphs(document):
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header,
                     section.first_page_footer, section.even_page_header,
                     section.even_page_footer):
            if part is None:
                continue
            yield from part.paragraphs
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def _get_notes_part(document, rel_type):
    """Find the footnotes or endnotes part via the main document part's
    relationships, and wrap each <w:footnote>/<w:endnote> paragraph as a
    python-docx Paragraph so we get the normal .runs/.text API.

    python-docx has no registered part class for footnotes/endnotes (as
    of 1.2.0), so when a docx is opened from disk this part loads as a
    generic opc.part.Part with no parsed `.element` -- we parse its raw
    bytes ourselves. When the part *was* constructed as an XmlPart in the
    same process (e.g. a part we just created), `.element` already
    exists and we reuse it directly instead of re-parsing.
    """
    main_part = document.part
    for rel in main_part.rels.values():
        if rel.reltype != rel_type or rel.is_external:
            continue
        notes_part = rel.target_part
        root = getattr(notes_part, "element", None)
        if root is None:
            # Loaded from disk as a generic (non-Xml) opc Part -- parsing
            # notes_part.blob gives us a *disconnected* tree, so any edits
            # we make wouldn't be picked up on save. Swap the relationship's
            # target for a proper XmlPart wrapping the parsed tree, so later
            # serialization reads from the (now editable, in-place) element.
            from docx.opc.part import XmlPart

            root = parse_xml(notes_part.blob)
            xml_part = XmlPart(notes_part.partname, notes_part.content_type, root, notes_part.package)
            rel._target = xml_part
            notes_part = xml_part
        # <w:footnote> / <w:endnote> elements, skipping the built-in
        # separator/continuationSeparator placeholders (type attr set).
        note_tag = "footnote" if "footnotes" in rel_type else "endnote"
        for note_el in root.findall(qn(f"w:{note_tag}")):
            note_type = note_el.get(qn("w:type"))
            if note_type in ("separator", "continuationSeparator"):
                continue
            for p_el in note_el.findall(qn("w:p")):
                yield Paragraph(p_el, note_el)
        return
    return


def iter_footnote_paragraphs(document):
    yield from _get_notes_part(document, FOOTNOTES_REL_TYPE)


def iter_endnote_paragraphs(document):
    yield from _get_notes_part(document, ENDNOTES_REL_TYPE)


def iter_all_paragraphs(document):
    """Everything: body, tables (incl. nested), headers/footers, footnotes, endnotes."""
    yield from iter_body_paragraphs(document)
    yield from iter_header_footer_paragraphs(document)
    yield from iter_footnote_paragraphs(document)
    yield from iter_endnote_paragraphs(document)


def count_all_paragraphs(document):
    return sum(1 for _ in iter_all_paragraphs(document))
