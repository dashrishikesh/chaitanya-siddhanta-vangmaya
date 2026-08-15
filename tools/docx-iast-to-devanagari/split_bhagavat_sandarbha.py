#!/usr/bin/env python3
"""Splits the Bhagavat-sandarbha source docx into two separate physical
docx files, BEFORE any transliteration -- same overall shape as
tattva-sandarbha's pipeline (split by commentator, then convert), except
tattva-sandarbha's split happens by color AFTER a single whole-document
convert.py pass, whereas here the split has to happen first: neither
color nor a fixed paragraph-index threshold reliably separates mula from
sarva-samvadini in this docx -- both color-based and idx-threshold-based
splitting were tried and both mis-attributed content (color: mula and
her prose are both default/black, since the same author writes both;
idx-threshold: tagged whole anucchedas as having her commentary when
they don't). The only reliable signal is her own explicit literal
"sarva-saṁvādinī:" label (occurs 7 times total in ~5200 paragraphs).

Each of the 102 anuccheda chunks (as delimited by the docx's own "[N]"
markers, see find_chunks) is built ENTIRELY from its own marker range --
no cross-chunk consolidation. Within that range, content BEFORE the
first "sarva-saṁvādinī:" label is mula; content FROM that label to the
end of the range is her voice. If no label falls within a given
anuccheda's own range, the entire chunk is mula -- no sarva-samvadini
content is invented or assumed for it.

Both files get a matching synthetic "[N]" marker at the start of any
chunk that has content for that side, so the two can be zipped back
together per-anuccheda after conversion (build_bhagavat_sandarbha_
site.py renders her portion as a "### सर्व-संवादिनी" subsection inside
that same anuccheda's section, not a separate trailing block).

Outputs (still IAST, ready for convert_no_hyphen_join.py):
  bhagavat_sandarbha_anucchedas.docx      -- mula content, per anuccheda
  bhagavat_sandarbha_sarva_samvadini.docx -- her voice, per anuccheda

Each paragraph is copied run-by-run, preserving original text AND color,
so color-based quote detection (for blockquote rendering) still works on
these split files exactly as it would on the original.
"""
import sys

from docx import Document

sys.path.insert(0, ".")
from docx_walk import iter_all_paragraphs
from paragraph_rewrite import paragraph_full_text
from generate_bhagavat_sandarbha import (
    DOCX_PATH,
    SARA_NISHKARSHA_START, SARA_NISHKARSHA_END,
    SARVA_SAMVADINI_ANUVYAKHYA_START, SARVA_SAMVADINI_ANUVYAKHYA_END,
    APPENDIX_START, APPENDIX_END,
    _SARVA_PREFIX_RE,
    find_chunks,
)

ANUCCHEDAS_OUT = "bhagavat_sandarbha_anucchedas.docx"
SARVA_SAMVADINI_OUT = "bhagavat_sandarbha_sarva_samvadini.docx"
SARA_NISHKARSHA_OUT = "bhagavat_sandarbha_sara_nishkarsha.docx"
SARVA_SAMVADINI_ANUVYAKHYA_OUT = "bhagavat_sandarbha_sarva_samvadini_anuvyakhya.docx"
APPENDIX_OUT = "bhagavat_sandarbha_appendix.docx"


def copy_paragraph(src_paragraph, dest_doc, override_text=None):
    """Copies a paragraph run-by-run, preserving text and color. If
    override_text is given (the "sarva-saṁvādinī:" label -- possibly with
    an inline "[N]" self-reference in front of it, see _SARVA_PREFIX_RE --
    already stripped off), the paragraph is written as a single run with
    that text instead; since the label is always uniformly formatted in
    this source, collapsing to one run here loses nothing.
    """
    dest_p = dest_doc.add_paragraph()
    if override_text is not None:
        dest_p.add_run(override_text)
        return dest_p
    for run in src_paragraph.runs:
        new_run = dest_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.color is not None and run.font.color.type is not None:
            try:
                new_run.font.color.rgb = run.font.color.rgb
            except Exception:
                pass
    return dest_p


def extract_flat_range(orig_paras, orig_para_objs, start, end, out_path):
    """Copies a flat idx range (no "[N]" markers, no mula/sarva split) to
    its own docx -- used for the standalone trailing sections (sara-
    nishkarsha, appendix) that sit outside any anuccheda's own range."""
    doc = Document()
    n = 0
    for idx in range(start, end):
        t = orig_paras[idx].strip()
        if not t or "o)0(o" in t:
            continue
        copy_paragraph(orig_para_objs[idx], doc, None)
        n += 1
    doc.save(out_path)
    return n


def main():
    src_doc = Document(DOCX_PATH)
    orig_para_objs = list(iter_all_paragraphs(src_doc))
    orig_paras = [paragraph_full_text(p).strip() for p in orig_para_objs]

    chunks, marker_indices = find_chunks(orig_paras)

    anucchedas_doc = Document()
    sarva_doc = Document()

    n_anuccheda_paras = 0
    n_sarva_paras = 0
    n_anuccheda_with_sarva = 0

    for label, start, end in chunks:
        mula_items = []   # (idx, override_text_or_None)
        sarva_items = []
        seen_label = False
        for idx in range(start, end):
            if idx in marker_indices:
                continue  # decimal sub-part markers (e.g. "[10.2]"), no content of their own
            t = orig_paras[idx].strip()
            if not t or "o)0(o" in t:
                continue
            m = _SARVA_PREFIX_RE.match(t)
            if m:
                seen_label = True
            override = t[m.end():] if m else None
            (sarva_items if seen_label else mula_items).append((idx, override))

        if label is not None:
            anucchedas_doc.add_paragraph().add_run(f"[{label}]")
            n_anuccheda_paras += 1
        for idx, override in mula_items:
            copy_paragraph(orig_para_objs[idx], anucchedas_doc, override)
            n_anuccheda_paras += 1

        if sarva_items:
            if label is not None:
                sarva_doc.add_paragraph().add_run(f"[{label}]")
                n_sarva_paras += 1
            n_anuccheda_with_sarva += 1
            for idx, override in sarva_items:
                copy_paragraph(orig_para_objs[idx], sarva_doc, override)
                n_sarva_paras += 1

    anucchedas_doc.save(ANUCCHEDAS_OUT)
    sarva_doc.save(SARVA_SAMVADINI_OUT)
    print(f"{ANUCCHEDAS_OUT}: {n_anuccheda_paras} paragraphs")
    print(f"{SARVA_SAMVADINI_OUT}: {n_sarva_paras} paragraphs "
          f"({n_anuccheda_with_sarva} chunks with sarva-samvadini content)")

    # The docx's own closing "sara-niskarsa" (essence-summary) section --
    # sits after mula's closing colophon and before her own separate
    # closing colophon, outside any numbered anuccheda's own range (see
    # generate_bhagavat_sandarbha.py). Extracted as its own file so it
    # becomes a standalone trailing site section, not duplicated inside
    # anuccheda 102.
    n_sara_paras = extract_flat_range(
        orig_paras, orig_para_objs, SARA_NISHKARSHA_START, SARA_NISHKARSHA_END, SARA_NISHKARSHA_OUT
    )
    print(f"{SARA_NISHKARSHA_OUT}: {n_sara_paras} paragraphs")

    # Right after sara-nishkarsha, the docx has its own two-line title
    # ("śrī-sarva-saṁvādinī" / "bhagavat-sandarbhānuvyākhyā") introducing a
    # further excerpt of her running commentary -- a distinct section, not
    # part of the summary above it.
    n_anuvyakhya_paras = extract_flat_range(
        orig_paras, orig_para_objs, SARVA_SAMVADINI_ANUVYAKHYA_START,
        SARVA_SAMVADINI_ANUVYAKHYA_END, SARVA_SAMVADINI_ANUVYAKHYA_OUT
    )
    print(f"{SARVA_SAMVADINI_ANUVYAKHYA_OUT}: {n_anuvyakhya_paras} paragraphs")

    # The docx's own "Appendix" section, right after her closing colophon --
    # a genuine supplementary discourse, not the editorial footnote
    # apparatus that follows it (excluded, see generate_bhagavat_sandarbha.py).
    n_appendix_paras = extract_flat_range(
        orig_paras, orig_para_objs, APPENDIX_START, APPENDIX_END, APPENDIX_OUT
    )
    print(f"{APPENDIX_OUT}: {n_appendix_paras} paragraphs")


if __name__ == "__main__":
    main()
