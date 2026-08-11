"""Builds a small synthetic .docx exercising every code path the
converter needs to handle: plain English (skip), IAST paragraphs,
mixed inline formatting, a table with IAST content, and a footnote.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_footnote(document, paragraph, footnote_text):
    """Minimal manual footnote insertion -- python-docx has no API for
    this, so we build the OOXML by hand (same relationship machinery the
    converter's docx_walk.py reads back out).
    """
    package = document.part.package
    main_part = document.part

    # Find or create the footnotes part.
    footnotes_part = None
    for rel in main_part.rels.values():
        if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes":
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        from docx.opc.part import XmlPart
        from docx.opc.packuri import PackURI
        from docx.oxml import parse_xml

        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            "</w:footnotes>"
        )
        partname = PackURI("/word/footnotes.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
        footnotes_part = XmlPart(partname, content_type, parse_xml(xml.encode("utf-8")), package)
        main_part.relate_to(
            footnotes_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )

    root = footnotes_part.element
    existing_ids = [int(fn.get(qn("w:id"))) for fn in root.findall(qn("w:footnote"))]
    new_id = max(existing_ids) + 1 if existing_ids else 1

    footnote_el = OxmlElement("w:footnote")
    footnote_el.set(qn("w:id"), str(new_id))
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = footnote_text
    r_el.append(t_el)
    p_el.append(r_el)
    footnote_el.append(p_el)
    root.append(footnote_el)

    # Reference run in the body paragraph.
    ref_run = paragraph.add_run()
    rpr = OxmlElement("w:rPr")
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rpr.append(vert)
    ref_run._element.append(rpr)
    ref_el = OxmlElement("w:footnoteReference")
    ref_el.set(qn("w:id"), str(new_id))
    ref_run._element.append(ref_el)


def main():
    document = Document()

    document.add_heading("Test Document", level=1)
    document.add_paragraph("This is a plain English paragraph that should be skipped entirely.")
    document.add_paragraph("Page 1 of 3")

    p1 = document.add_paragraph()
    p1.add_run("kṛṣṇaḥ sarva-kāraṇa-kāraṇam |")

    p2 = document.add_paragraph()
    p2.add_run("śrī-guruṁ ")
    bold_run = p2.add_run("caraṇāravinde")
    bold_run.bold = True
    p2.add_run(" nimittam apy uddhava ")
    # split a diacritic across two runs on purpose, to check the
    # run-concatenation logic actually protects against this:
    p2.add_run("kṛ")
    p2.add_run("ṣṇaḥ")

    p3 = document.add_paragraph()
    add_footnote(document, p3, "footnote body: nandīśvareśvara iti smṛtaḥ")
    p3.add_run("mūla-ślokas tathā vyākhyā ")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "yathā pāde nirṇītam"
    table.rows[0].cells[1].text = "plain english cell"

    document.save("test_input.docx")
    print("wrote test_input.docx")


if __name__ == "__main__":
    main()
