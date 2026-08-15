#!/usr/bin/env python3
"""Splits the Krishna-sandarbha source docx into separate physical IAST
docx files, before any transliteration -- see generate_krishna_
sandarbha.py's module docstring for the full methodology (same as
Bhagavat-sandarbha's and Paramatma-sandarbha's pipelines).

Outputs (still IAST, ready for convert_no_hyphen_join.py):
  krishna_sandarbha_anucchedas.docx      -- mula content, per anuccheda
  krishna_sandarbha_sarva_samvadini.docx -- her voice, per anuccheda

Unlike Bhagavat-sandarbha, this docx has no separate "sāra-niṣkarṣa"
essence-summary, no "Appendix" section, and no trailing sarva-saṁvādinī
excerpt -- her voice is fully integrated inline throughout (all 28 of
her explicit labels fall within the numbered anuccheda range).
"""
import sys

from docx import Document

sys.path.insert(0, ".")
from docx_walk import iter_all_paragraphs
from paragraph_rewrite import paragraph_full_text
from engines import looks_like_sanskrit, has_mixed_language_run
from generate_krishna_sandarbha import (
    DOCX_PATH,
    _SARVA_PREFIX_RE,
    find_chunks,
)

ANUCCHEDAS_OUT = "krishna_sandarbha_anucchedas.docx"
SARVA_SAMVADINI_OUT = "krishna_sandarbha_sarva_samvadini.docx"


def copy_paragraph(src_paragraph, dest_doc, override_text=None):
    """Copies a paragraph run-by-run, preserving text and color. If
    override_text is given (the "sarva-saṁvādinī:" label -- possibly with
    an inline "[N]"/"[N-M]" self-reference in front of it, see
    _SARVA_PREFIX_RE -- already stripped off), the paragraph is written as
    a single run with that text instead.
    """
    dest_p = dest_doc.add_paragraph()
    if override_text is not None:
        dest_p.add_run(override_text)
        return dest_p
    for run in src_paragraph.runs:
        new_run = dest_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font.color is not None and run.font.color.type is not None:
            try:
                new_run.font.color.rgb = run.font.color.rgb
            except Exception:
                pass
    return dest_p


def main():
    src_doc = Document(DOCX_PATH)
    orig_para_objs = list(iter_all_paragraphs(src_doc))
    orig_paras = [paragraph_full_text(p).strip() for p in orig_para_objs]

    chunks, marker_indices = find_chunks(orig_paras)

    anucchedas_doc = Document()
    sarva_doc = Document()

    n_anuccheda_paras = 0
    n_sarva_paras = 0
    n_anuccheda_with_sarva = 0

    for label, start, end in chunks:
        mula_items = []   # (idx, override_text_or_None)
        sarva_items = []
        seen_label = False
        for idx in range(start, end):
            if idx in marker_indices:
                continue  # decimal sub-part markers, no content of their own
            t = orig_paras[idx].strip()
            if not t or "o)0(o" in t:
                continue
            # This docx (unlike Bhagavat-sandarbha/Paramatma-sandarbha) has
            # inline editorial variant-reading notes mixed into the main
            # content flow, e.g. "[Vr. replaces above paragraph with the
            # following: ...]" -- English + IAST mixed. convert_no_hyphen_
            # join.py leaves paragraphs it can't convert UNCHANGED rather
            # than dropping them, so without this filter such notes would
            # leak into the site as raw Latin text. Drop anything that
            # wouldn't survive conversion anyway (same predicate the
            # converter itself uses).
            if not looks_like_sanskrit(t) or has_mixed_language_run(t):
                continue
            m = _SARVA_PREFIX_RE.match(t)
            if m:
                seen_label = True
            override = t[m.end():] if m else None
            (sarva_items if seen_label else mula_items).append((idx, override))

        if label is not None:
            anucchedas_doc.add_paragraph().add_run(f"[{label}]")
            n_anuccheda_paras += 1
        for idx, override in mula_items:
            copy_paragraph(orig_para_objs[idx], anucchedas_doc, override)
            n_anuccheda_paras += 1

        if sarva_items:
            if label is not None:
                sarva_doc.add_paragraph().add_run(f"[{label}]")
                n_sarva_paras += 1
            n_anuccheda_with_sarva += 1
            for idx, override in sarva_items:
                copy_paragraph(orig_para_objs[idx], sarva_doc, override)
                n_sarva_paras += 1

    anucchedas_doc.save(ANUCCHEDAS_OUT)
    sarva_doc.save(SARVA_SAMVADINI_OUT)
    print(f"{ANUCCHEDAS_OUT}: {n_anuccheda_paras} paragraphs")
    print(f"{SARVA_SAMVADINI_OUT}: {n_sarva_paras} paragraphs "
          f"({n_anuccheda_with_sarva} chunks with sarva-samvadini content)")


if __name__ == "__main__":
    main()
